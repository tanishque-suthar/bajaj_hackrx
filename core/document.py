import httpx
import fitz  # PyMuPDF
from docx import Document  # python-docx
import re
import logging
from typing import List, Optional, Tuple
from urllib.parse import urlparse
import tempfile
import os

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles PDF and DOCX document downloading, processing, and chunking"""
    
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        # Remove the synchronous client - we'll use async clients in methods
    
    async def process_document_from_url(self, url: str) -> List[str]:
        """
        Main method to process document from URL (supports PDF and DOCX)
        Returns list of text chunks
        """
        try:
            logger.info(f"Starting document processing for URL: {url}")
            
            # Step 1: Download document
            document_content = await self._download_document(url)
            
            # Step 2: Detect file type and extract text
            file_type = self._detect_file_type(url, document_content)
            raw_text = self._extract_text(document_content, file_type)
            
            # Step 3: Clean and preprocess
            cleaned_text = self._clean_and_preprocess(raw_text)
            
            # Step 4: Create chunks
            chunks = self._create_chunks(cleaned_text)
            
            logger.info(f"Document processed successfully. Created {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            raise Exception(f"Failed to process document: {str(e)}")
    
    async def _download_document(self, url: str) -> bytes:
        """Download document from URL (supports PDF and DOCX)"""
        try:
            logger.info("Downloading document from URL")
            
            # Validate URL
            parsed_url = urlparse(url)
            if not parsed_url.scheme or not parsed_url.netloc:
                raise ValueError("Invalid URL format")
            
            # Download with httpx async client
            async with httpx.AsyncClient(timeout=30.0) as client:
                response = await client.get(url)
                response.raise_for_status()
                
                # Log content type for debugging
                content_type = response.headers.get('content-type', '')
                logger.info(f"Downloaded document: {len(response.content)} bytes, content-type: {content_type}")
                
                return response.content
                
        except httpx.HTTPError as e:
            raise Exception(f"HTTP error downloading document: {str(e)}")
        except Exception as e:
            raise Exception(f"Error downloading document: {str(e)}")
    
    def _detect_file_type(self, url: str, content: bytes) -> str:
        """Detect file type from URL and content"""
        url_lower = url.lower()
        
        # Check URL extension first
        if url_lower.endswith('.pdf'):
            return 'pdf'
        elif url_lower.endswith('.docx'):
            return 'docx'
        
        # Check file signature (magic bytes)
        if content[:4] == b'%PDF':
            return 'pdf'
        elif content[:2] == b'PK' and b'word/' in content[:1000]:  # DOCX is a ZIP file containing word/ folder
            return 'docx'
        
        # Default to PDF for backward compatibility
        logger.warning("Could not determine file type, defaulting to PDF")
        return 'pdf'
    
    def _extract_text(self, content: bytes, file_type: str) -> str:
        """Extract text based on file type"""
        if file_type == 'pdf':
            return self._extract_text_from_pdf(content)
        elif file_type == 'docx':
            return self._extract_text_from_docx(content)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    def _extract_text_from_pdf(self, pdf_content: bytes) -> str:
        """Extract text from PDF using PyMuPDF"""
        try:
            logger.info("Extracting text from PDF")
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_file:
                temp_file.write(pdf_content)
                temp_file_path = temp_file.name
            
            try:
                # Open PDF with PyMuPDF
                doc = fitz.open(temp_file_path)
                text_content = []
                page_count = doc.page_count
                
                # Extract text from each page
                for page_num in range(page_count):
                    page = doc[page_num]
                    page_text = page.get_text()
                    
                    if page_text.strip():  # Only add non-empty pages
                        # Add page marker for context
                        text_content.append(f"\n--- Page {page_num + 1} ---\n")
                        text_content.append(page_text)
                
                # Combine all text
                full_text = "".join(text_content)
                
                # Close document and log results
                doc.close()
                logger.info(f"Extracted text from {page_count} pages, {len(full_text)} characters")
                
                return full_text
                
            finally:
                # Clean up temporary file
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                    
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")
    
    def _extract_text_from_docx(self, docx_content: bytes) -> str:
        """Extract text from DOCX using python-docx"""
        try:
            logger.info("Extracting text from DOCX")
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(docx_content)
                temp_file_path = temp_file.name
            
            try:
                # Open DOCX with python-docx
                doc = Document(temp_file_path)
                text_content = []
                
                # Extract text from paragraphs
                for para_num, paragraph in enumerate(doc.paragraphs):
                    para_text = paragraph.text.strip()
                    if para_text:  # Only add non-empty paragraphs
                        # Check if this looks like a heading (simple heuristic)
                        if paragraph.style.name.startswith('Heading') or (
                            len(para_text) < 100 and 
                            para_text.count('.') == 0 and 
                            para_text.isupper() == False and
                            para_text[0].isupper()
                        ):
                            text_content.append(f"\n--- {para_text} ---\n")
                        else:
                            text_content.append(para_text)
                
                # Extract text from tables
                for table_num, table in enumerate(doc.tables):
                    text_content.append(f"\n--- Table {table_num + 1} ---\n")
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            text_content.append(" | ".join(row_text))
                
                # Combine all text
                full_text = "\n".join(text_content)
                
                logger.info(f"Extracted text from {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables, {len(full_text)} characters")
                
                return full_text
                
            finally:
                # Clean up temporary file
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
                    
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")
    
    def _clean_and_preprocess(self, text: str) -> str:
        """Clean and preprocess extracted text"""
        try:
            logger.info("Cleaning and preprocessing text")
            
            # Remove excessive whitespace and normalize
            text = re.sub(r'\s+', ' ', text)
            
            # Remove special characters but keep important punctuation
            # Using a safer approach: remove only specific unwanted characters
            text = re.sub(r'[^\w\s\.,;:!?()\-\%\$\@\&\#\*\+\=\[\]\{\}<>/"\']', ' ', text)
            
            # Fix common OCR/extraction issues
            text = re.sub(r'\s+([.,;:!?])', r'\1', text)  # Remove space before punctuation
            text = re.sub(r'([.,;:!?])\s*([A-Z])', r'\1 \2', text)  # Ensure space after punctuation
            
            # Normalize quotes - replace Unicode quotes with standard ASCII
            text = re.sub(r'["\u201C\u201D]', '"', text)  # Replace various quote types
            text = re.sub(r"['\u2018\u2019]", "'", text)  # Replace various apostrophe types
            
            # Remove extra spaces
            text = re.sub(r'\s{2,}', ' ', text)
            
            # Remove page markers but keep structure (for PDF)
            text = re.sub(r'\n--- Page \d+ ---\n', '\n\n', text)
            
            # Remove table markers but keep structure (for DOCX)
            text = re.sub(r'\n--- Table \d+ ---\n', '\n\n', text)
            
            # Clean up heading markers - convert to paragraph breaks
            text = re.sub(r'\n--- (.*?) ---\n', r'\n\n\1\n\n', text)
            
            # Preserve paragraph breaks
            text = re.sub(r'\n\s*\n', '\n\n', text)
            
            # Clean up table separators
            text = re.sub(r'\s*\|\s*', ' | ', text)  # Normalize table separators
            
            # Trim and normalize
            text = text.strip()
            
            logger.info(f"Text cleaned: {len(text)} characters after preprocessing")
            return text
            
        except Exception as e:
            raise Exception(f"Error cleaning text: {str(e)}")
    
    def _create_chunks(self, text: str) -> List[str]:
        """Create overlapping text chunks"""
        try:
            logger.info(f"Creating chunks with size={self.chunk_size}, overlap={self.chunk_overlap}")
            
            # Split text into sentences first for better chunk boundaries
            sentences = self._split_into_sentences(text)
            
            chunks = []
            current_chunk = ""
            current_size = 0
            
            for sentence in sentences:
                sentence_size = len(sentence)
                
                # If adding this sentence would exceed chunk size
                if current_size + sentence_size > self.chunk_size and current_chunk:
                    # Add current chunk
                    chunks.append(current_chunk.strip())
                    
                    # Start new chunk with overlap
                    if self.chunk_overlap > 0:
                        overlap_text = self._get_overlap_text(current_chunk, self.chunk_overlap)
                        current_chunk = overlap_text + " " + sentence
                        current_size = len(current_chunk)
                    else:
                        current_chunk = sentence
                        current_size = sentence_size
                else:
                    # Add sentence to current chunk
                    if current_chunk:
                        current_chunk += " " + sentence
                    else:
                        current_chunk = sentence
                    current_size += sentence_size
            
            # Add final chunk if it exists
            if current_chunk.strip():
                chunks.append(current_chunk.strip())
            
            # Filter out very short chunks
            chunks = [chunk for chunk in chunks if len(chunk.strip()) > 50]
            
            logger.info(f"Created {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            raise Exception(f"Error creating chunks: {str(e)}")
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences using simple rules"""
        # Simple sentence splitting - can be improved with more sophisticated NLP
        sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', text)
        
        # Clean up sentences
        sentences = [sentence.strip() for sentence in sentences if sentence.strip()]
        
        return sentences
    
    def _get_overlap_text(self, text: str, overlap_size: int) -> str:
        """Get the last N characters for overlap"""
        if len(text) <= overlap_size:
            return text
        
        # Try to break at word boundary
        overlap_text = text[-overlap_size:]
        space_index = overlap_text.find(' ')
        
        if space_index > 0:
            return overlap_text[space_index:].strip()
        
        return overlap_text
    
    # Remove the __del__ method since we no longer have a persistent client

# Utility functions for external use
async def process_document(url: str) -> List[str]:
    """Convenience function to process a document (PDF or DOCX)"""
    processor = DocumentProcessor()
    return await processor.process_document_from_url(url)

def get_document_stats(chunks: List[str]) -> dict:
    """Get statistics about processed document chunks"""
    if not chunks:
        return {"total_chunks": 0, "total_characters": 0, "avg_chunk_size": 0}
    
    total_chars = sum(len(chunk) for chunk in chunks)
    avg_size = total_chars // len(chunks)
    
    return {
        "total_chunks": len(chunks),
        "total_characters": total_chars,
        "avg_chunk_size": avg_size,
        "min_chunk_size": min(len(chunk) for chunk in chunks),
        "max_chunk_size": max(len(chunk) for chunk in chunks)
    }
